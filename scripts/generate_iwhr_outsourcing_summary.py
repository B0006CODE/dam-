#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a filled (detailed) "科研外协合同工作总结报告" docx based on the IWHR template.

This script intentionally keeps values conservative and sourced from the provided
project close-out material (结项报告核心章节, 2026-02-02) plus the task description
in 项目结项报告.docx. Replace placeholders like period/unit/accept_date with the
actual contract info before submitting.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Pt


def _copy_paragraph_format(dst: Paragraph, src: Paragraph) -> None:
    # python-docx doesn't provide a direct clone; copy the most visible bits.
    try:
        dst.style = src.style
    except Exception:
        pass
    sf = src.paragraph_format
    df = dst.paragraph_format
    df.alignment = sf.alignment
    df.left_indent = sf.left_indent
    df.right_indent = sf.right_indent
    df.first_line_indent = sf.first_line_indent
    df.space_before = sf.space_before
    df.space_after = sf.space_after
    df.line_spacing = sf.line_spacing
    df.line_spacing_rule = sf.line_spacing_rule


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # type: ignore[attr-defined]
    new_para = Paragraph(new_p, paragraph._parent)  # type: ignore[arg-type]
    _copy_paragraph_format(new_para, paragraph)
    new_para.add_run(text)
    return new_para


def _format_table_compact(table) -> None:
    # Make tables readable: left align, fixed widths (set by caller), and remove
    # the "distributed" spacing effect inside narrow cells.
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0


def _set_table_col_widths(table, col_widths) -> None:
    # `col_widths` should be a list of Length objects (e.g. Cm/Pt/Inches).
    for idx, w in enumerate(col_widths):
        if idx < len(table.columns):
            table.columns[idx].width = w
    for row in table.rows:
        for idx, w in enumerate(col_widths):
            if idx >= len(row.cells):
                break
            row.cells[idx].width = w


def _insert_table_after_paragraph(paragraph: Paragraph, data: list[list[str]]):
    """Insert a table right after the given paragraph and return it."""
    # python-docx only appends tables; we create then move the XML node.
    doc = paragraph.part.document  # type: ignore[attr-defined]
    rows = len(data)
    cols = max((len(r) for r in data), default=0)
    if rows == 0 or cols == 0:
        return None

    table = doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            table.cell(r, c).text = data[r][c] if c < len(data[r]) else ""

    paragraph._p.addnext(table._tbl)  # type: ignore[attr-defined]
    _format_table_compact(table)
    return table


def _insert_table_after_table(doc: Document, table, data: list[list[str]]):
    """Insert a table right after the given table and return it."""
    rows = len(data)
    cols = max((len(r) for r in data), default=0)
    if rows == 0 or cols == 0:
        return None
    new_table = doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            new_table.cell(r, c).text = data[r][c] if c < len(data[r]) else ""
    table._tbl.addnext(new_table._tbl)  # type: ignore[attr-defined]
    _format_table_compact(new_table)
    return new_table


def _replace_paragraph_exact(doc: Document, placeholder: str, lines: list[str]) -> Paragraph:
    """
    Replace a paragraph whose text equals placeholder (after strip).
    Returns the first paragraph that was replaced (now containing lines[0]).
    """
    for para in doc.paragraphs:
        if (para.text or "").strip() == placeholder.strip():
            para.text = lines[0] if lines else ""
            cur = para
            for line in lines[1:]:
                cur = _insert_paragraph_after(cur, line)
            return para
    raise RuntimeError(f"placeholder not found: {placeholder[:40]}")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--template",
        type=Path,
        default=Path(__file__).resolve().parent.parent / ".." / "中国水利水电科学研究院科研外协合同工作总.docx",
        help="Path to IWHR template docx",
    )
    ap.add_argument(
        "--out",
        type=Path,
        default=Path(__file__).resolve().parent.parent / ".." / "科研外协合同工作总结报告_详细版.docx",
        help="Output docx path",
    )
    ap.add_argument("--contract-name", default="智能水利知识问答系统（Smart Water）基础支撑模块建设（知识图谱实例层标注/大模型部署环境/知识引擎基础界面）")
    ap.add_argument("--period", default="（待填：按合同起止时间填写）")
    ap.add_argument("--unit", default="（待填：外协单位全称）")
    ap.add_argument("--accept-date", default="（待填：拟申请验收日期）")
    args = ap.parse_args()

    doc = Document(args.template)

    # Fill the top summary table (4 rows, 2 cols).
    if not doc.tables:
        raise RuntimeError("template has no tables")
    tbl = doc.tables[0]
    tbl.rows[0].cells[1].text = args.contract_name
    tbl.rows[1].cells[1].text = args.period
    tbl.rows[2].cells[1].text = args.unit
    tbl.rows[3].cells[1].text = args.accept_date

    for para in doc.paragraphs:
        if (para.text or "").strip() == "（外协单位填写）":
            para.text = f"（外协单位填写：{args.unit}）"
            break

    # Section 1: Objectives / tasks / KPIs (detailed).
    obj_lines = [
        "（一）合同目标",
        "本外协合同聚焦于大坝安全业务领域知识图谱及知识引擎建设的基础性支撑环节，目标是在保障数据安全与系统稳定的前提下，完成数据与环境等基础支撑模块建设，为委托方后续核心模型研发与系统集成提供可靠支撑。",
        "说明：本外协工作不涉及主体算法模型研发与核心推理逻辑设计，相关内容由委托方总体方案统筹；外协单位按合同约定提供配套数据、环境与界面等工程支撑成果。",
        "",
        "（二）任务内容与主要考核指标（以合同及经委托方确认的交付清单为准）",
        "1. 知识图谱实例层基础标注与规范化处理",
        "1.1 工作内容：在委托方提供的概念层、本体/术语体系与数据源基础上，开展实例层数据的基础标注与结构化处理，包括实体识别、属性归类、关系标注、一致性校验与必要的规范化（同义归一、格式统一等）。",
        "1.2 交付物：实例层标注数据集；标注规范与质控说明文档（含字段/关系定义、标注流程、抽检方式与可追溯要求）。",
        "1.3 考核指标：标注数据准确性与一致性满足委托方抽检要求；数据可追溯（可定位来源与处理过程）；交付物格式与命名符合委托方入库/集成要求。",
        "",
        "2. 大模型部署所需硬件环境初始化配置与运行环境调优",
        "2.1 工作内容：按委托方部署架构与环境要求，完成服务器系统初始化、关键依赖安装、GPU/CPU 资源调度与基础性能测试；对容器编排与运行参数进行调优，保障训练/推理与检索/图谱等支撑服务可稳定运行。",
        "2.2 交付物：环境初始化配置方案；部署与运行说明文档（含组件启动方式、健康检查、备份建议、常见故障排查）。",
        "2.3 考核指标：按部署说明可复现启动；关键服务健康检查通过；能够支撑委托方开展训练/推理与知识库入库/查询等后续工作。",
        "",
        "3. 知识引擎人机交互基础界面（Web）开发",
        "3.1 工作内容：在委托方提供的接口与功能需求基础上，完成知识引擎基础交互界面原型与前端实现，覆盖基本信息展示、资料导入/任务进度、知识检索与结果展示等基础能力，并与委托方平台风格与接口协议保持一致，预留后续扩展接口。",
        "3.2 交付物：界面原型说明、前端代码与构建部署说明；接口对接说明与操作手册。",
        "3.3 考核指标：基础功能按需求可用；与后端接口联调通过；界面交互符合委托方平台规范，支持后续模块扩展与集成。",
        "",
        "4. 验收与配合事项",
        "4.1 工作内容：按委托方要求提供阶段汇报、问题响应、联调支持与验收材料整理。",
        "4.2 交付物：验收报告/交付清单、必要的测试与说明材料。",
        "",
        "（三）合同执行调整情况",
        "合同执行过程中如出现范围、指标或交付物调整，均以委托方书面确认/批准文件为准，并随验收材料归档。本报告所述内容与结项材料一致。",
    ]
    _replace_paragraph_exact(
        doc,
        "（应与外协合同内容一致，如在合同执行中有调整的，应说明调整情况并附相关批准文件）",
        obj_lines,
    )

    # Section 2.1: Completion evaluation (detailed).
    exec_lines = [
        "（一）总体执行情况",
        "外协单位按合同约定围绕“实例层标注、部署环境初始化、基础交互界面开发”三条主线推进实施，采用需求澄清-方案确认-开发实施-联调验证-交付验收的过程管理方式，按委托方要求完成交付与配合工作。",
        "",
        "（二）合同目标、任务完成情况（对应任务逐项说明）",
        "1. 知识图谱实例层基础标注与规范化处理",
        "1.1 实施过程：在委托方提供的概念层与数据源基础上，建立标注流程与质控机制，对实体/属性/关系进行基础标注与一致性校验；对关键字段执行规范化处理，确保同一术语/实体的表达统一，降低后续入库与检索噪声。",
        "1.2 质量控制：采用抽检与一致性校验相结合的方式进行质控；标注结果可追溯到原始资料来源与处理记录，便于复核与审计。",
        "1.3 完成情况：形成标注数据集及标注规范/质控说明文档，并按委托方要求完成格式整理与交付。",
        "",
        "2. 大模型部署所需硬件环境初始化配置与运行环境调优",
        "2.1 实施过程：完成服务器系统初始化与关键依赖安装，配置容器运行环境与必要的网络/权限策略；对 GPU/CPU 资源调度与基础性能进行验证，保障部署可用与稳定运行。",
        "2.2 可复现与运维：形成部署与运行说明，包含组件启动方式、健康检查、备份范围建议与常见故障排查要点，便于委托方运维接手与后续扩展。",
        "2.3 完成情况：完成运行环境初始化配置与调优，并按要求提交环境配置方案与说明文档。",
        "",
        "3. 知识引擎人机交互基础界面（Web）开发",
        "3.1 实施过程：基于委托方接口与功能需求，完成基础界面原型与前端实现，覆盖资料导入、任务中心、检索查询与结果展示等基础交互；界面风格与交互逻辑与委托方平台保持一致。",
        "3.2 联调验证：完成与后端接口的联调验证，支持流式输出/查询展示等基础能力，并预留后续智能推理与服务模块的接入点。",
        "3.3 完成情况：交付界面原型与前端代码（含构建部署说明、接口对接说明与操作手册）。",
        "",
        "4. 交付与验收配合",
        "按委托方验收要求整理交付清单与材料，配合开展问题响应、联调验证与验收准备工作。",
    ]
    _replace_paragraph_exact(doc, "（对应合同逐一说明完成情况）", exec_lines)

    # Section 2.2: Achievements (detailed + snapshots).
    ach_lines = [
        "（一）主要成果与交付物",
        "1) 实例层标注数据集及标注规范/质控说明文档。",
        "2) 部署环境初始化配置方案与运行说明文档（含启动、健康检查、备份建议与排障要点）。",
        "3) 知识引擎基础交互界面原型与前端代码（含接口说明与操作手册）。",
        "4) 验收材料与交付清单（配合委托方验收）。",
        "",
        "（二）解决的关键工程技术问题（与委托方总体方案配套）",
        "1) 多源资料接入与结构化处理链路的工程化支撑：支持 PDF/图片/Word/网页/表格等多格式资料进入统一处理流程；对扫描件资料可通过 OCR（如 MinerU）转为结构化 Markdown，并保留标题层级与表格表达，确保后续分块、检索与证据引用的一致性。",
        "2) 可追溯与可控的工程实现：为文件与分块附带来源与元信息；支持内容哈希（如 SHA-256）用于去重与增量导入；记录入库状态（done/failed）便于重试与问题定位；对路径与输入进行校验降低安全风险。检索侧可按阈值、top_k、重排序等策略控制召回边界（按委托方配置与策略执行）。",
        "3) 内网可落地的部署与运维支撑：以容器编排方式组织关键组件，支持内网离线镜像打包与迁移；提供备份范围建议（如容器卷数据、配置与审计库等）与健康检查要点，降低交付后的运维成本。",
        "4) 可替换与可扩展的工程设计配合：对话模型、Embedding、OCR、向量库、重排序器等组件以配置方式解耦，便于委托方按资源与安全要求替换；基础界面预留后续推理与智能服务模块的接入点。",
        "",
        "（三）关键数据与系统快照（源自结项材料，便于验收复核）",
        "1. 结构化基础数据规模：水库基础信息数据共 1512 条记录（用于统计、检索过滤与地图展示）。",
        "2. 图谱示例系统快照：实体数 11788、关系数 25474（更新时间 2026-01-15）。",
        "3. 支撑组件版本快照（以 docker-compose 配置为准）：Neo4j 5.26、Milvus v2.5.6、MinIO RELEASE.2023-03-20T20-16-18Z、etcd v3.5.5（MinerU 可选）。",
        "",
        "（四）应用情况（面向水库/大坝工程领域）",
        "形成面向工程问答、资料检索与归纳、统计洞察与图谱关系查询的基础能力，为委托方后续业务集成、场景化应用与持续迭代提供支撑。",
        "",
        "（五）知识产权与成果权益",
        "本外协工作形成的数据、文档与软件代码等成果的权属与使用范围，按外协合同约定执行；如需申请专利、软著或标准等，配合委托方提供必要材料与证明。",
        "补充说明：本外协工作以工程支撑与交付为主，未单独开展国际/国内同类产品指标对比测试；如委托方需要，可在后续阶段按统一评测口径组织对比与第三方测试。",
    ]
    ach_anchor = _replace_paragraph_exact(
        doc,
        "（解决的关键技术问题、取得的重大科技成果，与国际国内同类技术或产品的技术、经济和环保指标的比较；项目在水利生产实践中的应用情况；取得专利、软件著作权等知识产权情况，技术标准情况，以及相关知识产权和成果的权益归属；建成的试验基地、生产线、示范点等情况；人才队伍建设情况等。）",
        ach_lines,
    )

    # Insert supporting tables after the "(三)关键数据与系统快照..." block to make it more concrete.
    # We find the paragraph containing that heading text.
    snap_para: Paragraph | None = None
    for p in doc.paragraphs:
        if (p.text or "").strip() == "（三）关键数据与系统快照（源自结项材料，便于验收复核）":
            snap_para = p
            break

    if snap_para is not None:
        section = doc.sections[0]
        content_width = int(section.page_width - section.left_margin - section.right_margin)
        # 3-col table: give enough width to the "用途" column to avoid odd wrapping.
        col1 = int(content_width * 0.18)
        col2 = int(content_width * 0.44)
        col3 = content_width - col1 - col2
        versions_tbl = _insert_table_after_paragraph(
            snap_para,
            [
                ["组件", "用途", "版本/镜像"],
                ["Neo4j", "知识图谱存储与查询", "neo4j:5.26"],
                ["Milvus", "向量检索", "milvusdb/milvus:v2.5.6"],
                ["MinIO", "对象存储（Milvus 依赖/文件资源）", "minio/minio:RELEASE.2023-03-20T20-16-18Z"],
                ["etcd", "Milvus 元数据依赖", "quay.io/coreos/etcd:v3.5.5"],
                ["MinerU（可选）", "PDF/图片 OCR 与版面理解", "mineru-sglang:latest（本地构建）"],
            ],
        )
        if versions_tbl is not None:
            _set_table_col_widths(versions_tbl, [Emu(col1), Emu(col2), Emu(col3)])
            stats_tbl = _insert_table_after_table(
                doc,
                versions_tbl,
                [
                    ["统计项（基于 1512 条结构化数据）", "结果"],
                    ["坝高分级（maxHeight）", "低坝(<30m) 839(55.67%)；中坝(30-70m) 637(42.27%)；高坝(70-100m) 28(1.86%)；特高坝(>=100m) 3(0.20%)"],
                    ["坝型（汇总口径）", "土石坝 1308(86.51%)；混凝土坝 122(8.07%)；拱坝 31(2.05%)；其他/未分类 51(3.37%)"],
                    ["规模分布", "中型 1344(88.89%)；大型 164(10.85%)；未知 4(0.26%)"],
                    ["流域分布（Top）", "长江 644(42.59%)；珠江 261(17.26%)；淮河 180(11.90%)；黄河 134(8.86%)；松辽 123(8.13%)"],
                ],
            )
            # 2-col table: first col short label, second col wide result.
            if stats_tbl is not None:
                c1 = int(content_width * 0.28)
                c2 = content_width - c1
                _set_table_col_widths(stats_tbl, [Emu(c1), Emu(c2)])

    # Funding section: expand with a (placeholder) breakdown.
    fund_lines = [
        "（一）经费使用原则",
        "本外协经费使用严格按照合同约定与财务制度执行，坚持专款专用、凭证齐全、过程可追溯的原则；涉及采购/服务等事项按委托方及外协单位内部制度履行审批与留痕。",
        "",
        "（二）经费主要支出方向（示例分类，最终以财务科目与凭证为准）",
        "1) 人工与技术服务投入：数据标注、前后端开发、部署与调优、测试与文档编制等。",
        "2) 环境与算力相关支出：服务器环境配置、必要的软件工具与部署支撑等。",
        "3) 项目管理与验收支撑：阶段汇报、问题响应、联调支持、验收材料整理等。",
        "",
        "（三）经费明细与附件",
        "经费明细表、合同约定的付款节点材料及相关财务凭证，可随验收附件一并提交（如委托方要求）。",
        "（待填：合同总经费/本期已使用金额/科目明细）",
    ]
    # Insert after the heading, without deleting the heading itself.
    for para in doc.paragraphs:
        if (para.text or "").strip() == "经费使用情况":
            cur = para
            for line in fund_lines:
                cur = _insert_paragraph_after(cur, line)
            break

    # Management experience: add more detail.
    mgmt_lines = [
        "（一）组织与分工",
        "围绕“数据标注/环境部署/界面开发/文档与验收”设置责任人，建立日常沟通与问题响应机制；关键事项形成会议纪要与确认记录，确保需求、接口与交付口径一致。",
        "",
        "（二）计划与过程管理",
        "采用里程碑管理方式推进：需求澄清与方案确认 → 阶段开发实施 → 联调验证 → 交付与验收材料整理；对需求变更、问题单与风险项建立台账，按优先级闭环处理。",
        "",
        "（三）质量与安全管理",
        "1) 数据质量：对标注数据执行一致性校验与抽检复核，交付前按委托方要求进行格式检查与可追溯性核对。",
        "2) 工程质量：关键功能点进行联调与回归验证；部署配置与运行说明同步维护，确保可复现与可交接。",
        "3) 安全合规：资料处理与环境部署遵循委托方数据安全要求，采用最小权限与必要隔离策略；交付材料按版本管理留痕，便于审计。",
        "",
        "（四）文档与交付管理",
        "对交付物建立清单化管理，包含版本号、日期、责任人与变更记录；验收阶段按委托方模板与要求整理材料并提供必要说明。",
    ]
    for para in doc.paragraphs:
        if (para.text or "").strip() == "外协单位对本外协合同的组织管理经验":
            cur = para
            for line in mgmt_lines:
                cur = _insert_paragraph_after(cur, line)
            break

    # Issues & suggestions: expand.
    issue_lines = [
        "（一）存在问题",
        "1) 原始资料质量与格式差异较大（扫描件、表格、图片等），OCR 与结构化抽取存在一定误差，需持续优化资料清洗规范并建立抽检机制。",
        "2) 领域术语与关系定义需要在实际应用中持续沉淀，同义归一与关系类型边界需结合业务复核不断迭代。",
        "3) 内网部署环境差异（硬件、驱动、网络策略）可能带来迁移成本，需在交付阶段提供更细化的环境清单与排障手册。",
        "",
        "（二）建议",
        "1) 建立“增量入库 + 质量评分 + 可重试”机制，对解析失败或质量较低资料形成闭环改进流程。",
        "2) 建立评测与反馈闭环：沉淀典型问题集与标准化问法模板，结合用户反馈持续优化检索策略与提示词/工具编排策略。",
        "3) 推进权限与审计策略落地：细化知识库/图谱/模型配置权限分级，完善日志与审计留痕，满足单位安全要求。",
        "4) 推动版本化管理：对知识库/图谱/配置进行版本快照与差异比对，支持回滚与复盘，降低迭代风险。",
    ]
    for para in doc.paragraphs:
        if (para.text or "").strip() == "存在问题及建议":
            cur = para
            for line in issue_lines:
                cur = _insert_paragraph_after(cur, line)
            break

    # Attachments: make it more explicit.
    att_lines = [
        "1. 项目结项报告（核心章节，版本 V1.0，日期 2026-02-02）。",
        "2. 实例层标注数据集（含数据说明/字段字典，如有）。",
        "3. 标注规范与质控说明文档。",
        "4. 部署环境初始化配置方案与运行说明文档（含健康检查、备份与排障要点）。",
        "5. 知识引擎基础交互界面原型说明、前端代码与构建部署说明。",
        "6. 接口对接说明与操作手册。",
        "7. 验收材料与交付清单。",
        "8. 经费明细与相关财务凭证（如委托方要求）。",
    ]
    for para in doc.paragraphs:
        if (para.text or "").strip() == "附件":
            cur = para
            for line in att_lines:
                cur = _insert_paragraph_after(cur, line)
            break

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
