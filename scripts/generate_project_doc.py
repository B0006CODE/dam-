from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "智能水利知识问答系统（Smart Water）项目说明书.docx"


def _set_run_font(
    run,
    *,
    ascii_font: str | None = None,
    east_asia_font: str | None = None,
    size_pt: int | None = None,
) -> None:
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    if ascii_font:
        rfonts.set(qn("w:ascii"), ascii_font)
        rfonts.set(qn("w:hAnsi"), ascii_font)
    if east_asia_font:
        rfonts.set(qn("w:eastAsia"), east_asia_font)
    if size_pt:
        run.font.size = Pt(size_pt)


def _set_document_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style.element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    for section in document.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)


def _add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def _add_toc(document: Document) -> None:
    document.add_heading("目录", level=1)
    p = document.add_paragraph()
    _add_field(p, 'TOC \\o "1-2" \\h \\z \\u')
    tip = document.add_paragraph("提示：打开 Word 后，右键目录区域 →“更新域”→“更新整个目录”。")
    tip.runs[0].italic = True


def _add_cover(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("智能水利知识问答系统（Smart Water / Yuxi‑Know）\n项目说明书")
    _set_run_font(run, ascii_font="Calibri", east_asia_font="黑体", size_pt=26)
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("——让资料变得可检索、让回答变得可追溯")
    _set_run_font(run, ascii_font="Calibri", east_asia_font="楷体", size_pt=14)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = date.today().isoformat()
    run = meta.add_run(f"\n版本：v1.0\n日期：{today}\n")
    _set_run_font(run, ascii_font="Calibri", east_asia_font="宋体", size_pt=12)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("（说明：本文档以叙述方式讲清楚“做了什么、为什么这样做、怎么落地运行”。）")
    _set_run_font(run, ascii_font="Calibri", east_asia_font="宋体", size_pt=11)
    run.italic = True


def _p(document: Document, text: str) -> None:
    for line in text.split("\n"):
        document.add_paragraph(line.rstrip())


def _table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def _count_chars(texts: list[str]) -> int:
    raw = "".join(texts)
    raw = re.sub(r"\s+", "", raw)
    return len(raw)


def build_document() -> tuple[Document, int]:
    document = Document()
    _set_document_defaults(document)

    all_texts: list[str] = []

    _add_cover(document)
    document.add_page_break()

    _add_toc(document)
    document.add_page_break()

    document.add_heading("前言：把系统写成“能被使用”的样子", level=1)
    preface = (
        "这套系统的出发点并不浪漫：它解决的是一个很现实的场景——资料越来越多，但在关键时刻反而更难用。\n"
        "规范、规程、设计说明、验收记录、运行台账、巡检报告、监测异常清单、历史处置方案……这些东西都很重要，"
        "但它们分散、版本多、格式杂。真正要用的时候，很多单位会回到一种“熟人系统”：谁记得在哪、谁当年做过、谁说了算。\n\n"
        "我们希望把这种“靠人兜底”的方式，变成一种“靠系统也站得住”的能力。为此，我们没有把大模型当成万能专家，"
        "而是把它当成一个助理：它负责理解提问、组织表达、把结论写得更清楚；但它必须围绕证据工作，证据来自你自己的资料。\n\n"
        "这也是本文档的写法：尽量用连续叙述把脉络讲清楚，而不是用“点点点”把读者淹没。"
        "你读完应该能回答三个问题：它解决什么，它怎么运转，它怎么在内网和日常运维里长期跑下去。\n\n"
        "如果你是使用者，可以重点看第一、二、六章；如果你负责部署运维，重点看第五章；如果你要二次开发或接入业务系统，"
        "第四章会告诉你代码大致在哪里、结构是怎么划分的。"
    )
    _p(document, preface)
    all_texts.append(preface)

    document.add_heading("第一章 它解决什么问题（以及它不解决什么）", level=1)
    t = (
        "最常见的三连是：找不到、看不懂、用不上。\n"
        "资料散落在文件夹、网盘、OA、群聊里；同一个问题可能有多个版本答案；关键结论藏在扫描 PDF 里；新同事靠口口相传。"
        "一旦遇到紧急情况（异常监测、汛期调度、巡检发现隐患），时间窗口短，靠人工翻资料既慢又容易遗漏。\n\n"
        "你可能见过这样的画面：值班人员面对一条异常点位，脑子里第一反应不是“我能不能问系统”，而是“我该找谁”。"
        "找人当然有效，但它依赖的是个人记忆和即时响应；而工程管理更需要的是可重复、可审计的流程。\n\n"
        "系统的目标很明确：把资料变成可检索资产，让回答尽量可追溯，把重复劳动交给系统，把判断与决策留给专家。"
        "它不会替你做最终签字，也不会承诺“问什么都能答对”。相反，它更像一个可靠的助理：它能把最相关的依据找出来，"
        "把依据说清楚，并在资料不足时坦诚地告诉你“我没在库里找到”，同时建议下一步要补什么资料、要怎么追问。\n\n"
        "这也是它“有用但不夸张”的地方：如果你把它当作搜索引擎，它比搜索引擎更懂语义；如果你把它当作聊天机器人，它比聊天机器人更在意出处。"
        "而当它开始被真正使用时，它也会反过来暴露一个事实：资料本身哪里不完整、哪里不统一、哪里需要治理。"
    )
    _p(document, t)
    all_texts.append(t)

    document.add_heading("第二章 它怎么运转：从资料到可信回答", level=1)
    t = (
        "系统的工作方式可以概括成一句话：先把资料变成可检索的资产，再让模型围绕这些资产回答问题。\n\n"
        "资料进入系统之后，不会直接被“丢给模型去读”。系统会先做统一处理：不管是 Word、PDF、Excel 还是图片，"
        "尽量转成结构更统一的文本（通常是 Markdown 或纯文本）。这一步的意义很直观："
        "Word 有标题层级，PDF 有分页版面，Excel 更像表格；如果直接混着切块，边界会很随机。"
        "统一成 Markdown 后，标题、列表、段落就有了可依赖的结构，切块也更像“按逻辑切开”，而不是“按字符硬切”。\n\n"
        "每个片段都会带上元数据：它来自哪个文件、在文件里的顺序、对应的 file_id 等。它看起来琐碎，但它决定了后面能不能追溯、"
        "能不能清理、能不能导出迁移。很多系统后期做不下去，不是模型问题，而是最开始没有把“可治理”留在数据结构里。\n\n"
        "切块本身也不是越细越好。太大，会把不相关的内容一起带回来，模型读起来噪声大；太小，又容易把一句话切成两半，"
        "导致答案缺一截。系统采用相对稳妥的默认值（例如 chunk_size、overlap），并保留可配置空间，"
        "让你在“召回率”和“可读性”之间做权衡。某些资料本身是一问一答式的，系统也支持按 QA 结构切分，"
        "用更自然的边界减少‘切碎语义’。\n\n"
        "接下来是向量化（Embedding）：系统把每个片段变成一串向量坐标，写进向量库（默认 Milvus，也保留轻量的 Chroma 作为备选）。"
        "当用户提问时，系统同样把问题转成向量，在向量库里找出最相近的片段，把它们作为证据交给模型。"
        "模型在证据范围内组织答案，而不是凭空编造一个看起来很合理的结论。\n\n"
        "在实际效果上，影响最大的往往不是“有没有检索”，而是“检索回来的东西能不能直接用”。"
        "工程资料里经常出现一种情况：段落看起来相关，但恰好缺了关键条件或例外条款，于是回答就会变得含糊。"
        "所以系统还支持重排序（Re-ranker）：先粗召回一批候选片段，再用更精细的模型把更能回答问题的片段排到前面。"
        "这一步看似细节，但在规范条文、条件判断这类问题上，经常是体验分水岭。\n\n"
        "有些问题并不是在问“解释”，而是在问“关系”或“统计”。比如“某类缺陷最常见的原因是什么”“某个实体关联了哪些处置措施”。"
        "这类问题如果只靠向量检索，常常会变成“给你一堆段落，你自己总结”。所以系统还提供知识图谱路径："
        "通过 Neo4j + LightRAG 抽取实体与关系，支持子图查询、扩展、统计，让系统能把关系链讲清楚。\n\n"
        "资料里还有一个顽固敌人：扫描件。图片清晰，但文本不可复制；很多关键资料因此天然“不可检索”。"
        "系统提供 OCR 插件，默认路线更推荐 MinerU（在 docker-compose 中以 profile=all 方式启用，适合有 GPU 的场景）。"
        "另外也保留 PaddleX 作为可选方案（当前 compose 中默认注释），便于不同环境取舍。\n\n"
        "这里有一个很实用的经验：OCR 不一定越多越好。对可直接解析出文本的 PDF，优先走文本解析路线，"
        "避免 OCR 引入误差；对扫描件，再用 OCR 补齐。系统提供 OCR 健康检查与统计，是为了让这件事“可观察”："
        "你能知道服务是否可用、失败率是否异常，而不是只在用户抱怨时才发现问题。\n\n"
        "最后还有一个很现实的问题：导入、OCR、向量化往往是分钟级工作，如果都堵在接口里，体验和稳定性都会崩。"
        "所以系统把这些工作放到后台任务队列里，前端只显示进度；对话则支持流式输出，用户不会只看一个加载转圈。"
        "更重要的是，流结束后系统会把会话、消息、工具调用落库——这让“这句话怎么来的”变得可追溯，而不是只剩下一个最终答案。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "你可以把后台任务理解成“把重活挪到后台做”。"
        "导入时系统会不断更新任务进度：正在解析、正在 OCR、正在向量化、正在入库。"
        "这件事对用户体验的意义是直观的：你不会因为导入一个大文件就把页面卡死；对运维的意义也很直观："
        "当导入失败时，你能看到失败发生在哪一步，而不是只得到一句“失败了”。"
        "任务也支持取消与清理，是为了避免在错误参数、错误文件、或资源不足时把系统拖进持续消耗的状态。\n\n"
        "对话侧的落库同样是为了让系统可运营。"
        "在工程场景里，很多改进并不是靠“灵感”，而是靠复盘：为什么这次回答偏了？当时检索到了什么？"
        "工具调用发生了什么？用户是怎么追问的？这些信息如果不落库，系统就永远只能靠主观感觉迭代。"
    )
    _p(document, t)
    all_texts.append(t)

    flow = (
        "资料导入（文件/URL）→ 解析/清洗 → 统一为 Markdown → 分块（chunk）→ Embedding → 写入向量库（Milvus/Chroma）\n"
        "                                    └（可选）LightRAG 抽取实体/关系 → 写入 Neo4j（图谱）\n"
        "用户提问 → Agent（LangGraph）选择工具（检索/图谱/统计）→ 汇总证据 → 组织回答（流式返回）→ 会话与工具调用落库"
    )
    p = document.add_paragraph()
    run = p.add_run(flow)
    _set_run_font(run, ascii_font="Consolas", east_asia_font="等线", size_pt=10)
    all_texts.append(flow)

    t = (
        "为了让你更直观地理解这套机制，我们可以用一条真实的“问答旅程”来想象它。\n\n"
        "用户问：“混凝土面板出现裂缝后，通常需要关注哪些风险？先做哪些处置？”系统不会立刻让模型自由发挥。"
        "它会先判断这类问题更像“需要依据的解释型问题”，于是去知识库里检索“裂缝、面板、处置、监测、渗漏”等相关片段，"
        "必要时再做重排序，把包含关键条件的段落往前提。"
        "如果用户接着追问“有没有类似案例？常见原因排名是什么？”这时系统更可能触发图谱路径，用子图和统计把“原因—现象—措施”的结构化关系讲清楚。\n\n"
        "对用户来说，这些动作不是为了“显得高级”，而是为了让回答更像一个工程人员会写的分析：有依据，有条件，有边界，"
        "并且能被下一位同事拿去复核和继续追问。"
    )
    _p(document, t)
    all_texts.append(t)

    document.add_heading("第三章 原理但不枯燥：RAG、图谱与 Agent", level=1)
    t = (
        "很多人第一次接触这类系统，会问：既然有大模型，为什么还要搞数据库、搞图谱、搞流程？\n"
        "答案是：工程场景要的是“可控”和“可追溯”。模型的表达能力很强，但它天生不擅长对内部资料负责。\n\n"
        "RAG 的核心价值，是把模型的回答范围钉在你的资料上。你可以把它理解成：先从资料柜里把相关页码抽出来，再让模型做归纳。"
        "这样做并不神秘，也不需要重新训练模型：资料更新后立刻生效，回答也更容易做到“有据可查”。\n\n"
        "Embedding 可以理解成给文本装上一个“语义坐标”。问题和资料片段都变成向量后，就能做相似度检索。"
        "但相似不等于能回答，所以系统支持重排序（Re-ranker）：先粗召回一批，再精排，把更“能用”的片段往前提。"
        "工程场景里，很多回答差的原因不是没检索到，而是检索到的片段刚好缺了关键条件或例外条款。\n\n"
        "一个更贴近直觉的例子是：你问“裂缝出现后要不要停机处理”。"
        "资料里可能有很多段落同时出现“裂缝”这个词，但真正决定动作的往往是裂缝位置、宽度、发展趋势、结构类型等条件。"
        "重排序的目的，就是把那些包含关键条件的段落放到更前面，让模型更容易做出“有条件、有边界”的回答，而不是泛泛而谈。\n\n"
        "图谱适合解决另一类问题：不是找段落，而是找关系。向量检索擅长回答“这是什么、怎么理解”，"
        "图谱更擅长回答“有哪些、多少、排名、分布、关联路径”。LightRAG 的作用是把图谱抽取与问答结合起来："
        "导入时抽实体关系，查询时拿子图，最后由模型把子图解释成更易读的结论。\n\n"
        "如果你把图谱界面打开，会更容易理解它的价值：它不是为了炫技，而是为了让“为什么这么说”变得可视化。"
        "当系统建议先排查某个原因时，你可以沿着图谱链路看到它是如何关联到缺陷、措施、案例的。"
        "这种“看得见的逻辑”对培训、复核、验收都很重要。\n\n"
        "Agent（LangGraph）把对话从一次调用变成一个流程。现实里的问答很少一步到位：你会追问、会缩小范围、会要求给依据、"
        "会要求结构化输出。用状态机编排步骤的好处是：流程可解释、可插拔、可回放；当你需要更严格的边界控制时，"
        "也更容易落在流程里，而不是靠提示词碰运气。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "还有一个常见的落地问题是：模型到底放在哪里？\n\n"
        "系统把模型当成“可替换的能力”，而不是写死在代码里。你既可以用云端模型（比如单位已有采购或可出网环境），"
        "也可以用本地模型（内网、成本敏感、或对数据出域有严格要求）。"
        "本地路线上，Ollama 更像“一键能跑”的选择，适合快速验证；vLLM 更偏生产，适合 GPU 服务器与并发。"
        "系统用统一的模型配置方式把这些入口收敛起来：对前端来说是在列表里选模型，对后端来说是根据 base_url 和 api_key 去调用对应服务。\n\n"
        "这种“模型可替换”的设计很重要，因为你永远不知道单位的网络、合规、预算会怎么变。系统把知识留在自己的存储里，"
        "把模型当成外部能力去接入，这样才有长期演进的空间。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "在一些单位环境里，还会关心“回答边界”：哪些内容可以说，哪些内容必须谨慎。\n\n"
        "系统提供了内容安全相关的开关与策略接口，便于在需要时加上护栏；但更关键的仍然是“证据约束”。"
        "让回答尽量落在检索证据上，让不确定就明确说明，这比任何花哨的措辞都更接近工程场景的可信。\n\n"
        "与其追求一次回答写得多漂亮，不如把“引用、边界、追问引导”这三件事做稳。"
        "当系统能稳定地把人带回到资料与证据上，它就已经在工程管理里发挥作用了。"
    )
    _p(document, t)
    all_texts.append(t)

    document.add_heading("第四章 系统落地：你能在项目里找到的关键部件", level=1)
    t = (
        "如果你只是使用系统，这一章可以快速浏览；但如果你要维护或二次开发，知道“代码大概在哪”会省很多时间。\n\n"
        "后端入口在 server/main.py，基于 FastAPI 对外提供统一 API。路由按业务拆分在 server/routers 下，"
        "包括认证、对话、知识库、图谱、系统状态、任务中心等。对话接口支持流式返回，并在流结束后把 LangGraph 状态写回数据库，"
        "让对话过程可追溯。\n\n"
        "认证侧既支持账号登录，也提供 SSO 登录接口，便于接入单位现有门户。安全策略上，登录接口有频率限制，用户模型也记录失败次数与锁定时间，"
        "目的是减少暴力破解暴露面。上传侧则做了路径校验，避免用路径遍历读取不该读取的文件。\n\n"
        "知识库实现主要在 src/knowledge：抽象基类定义统一接口，不同实现（Milvus/Chroma/LightRAG）放在 implementations。"
        "导入过程里包含解析、OCR、转 Markdown、切块、向量化、入库；OCR 通过 src/plugins/_ocr.py 封装，"
        "MinerU 与 PaddleX 作为服务接入。\n\n"
        "存储层分两块：一块是对话与审计（SQLite + SQLAlchemy，Conversation/Message/ToolCall 等模型），"
        "另一块是对象存储（MinIO），用于头像、图片等二进制数据。对话落库的意义不是“存聊天记录”这么简单，"
        "它更像是把系统的工作过程沉淀成可运营数据：哪些问题高频、哪些工具常用、哪些工具失败多、哪些库使用最活跃，都可以被看见。\n\n"
        "前端在 web 目录，是 Vue3 + Pinia + vue-router 的结构，主要负责交互与可视化（包括图谱展示）。"
        "从使用体验上，它把“导入任务进度”“对话流式输出”“图谱浏览”“系统配置”这些动作统一在一个界面里，"
        "让系统不只是一个 API，而是能被非开发人员日常使用的产品。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "项目里还专门放了一类“贴近业务的数据服务”，例如大坝监测异常数据的拉取与整理。"
        "这种能力不是为了把系统变成监测平台，而是为了把“业务数据”变成问答可用的上下文："
        "当你面对一串异常点位和评分时，系统可以先把它们整理成一段更易读的摘要，再引导你提问‘我该先排查哪里、该看哪些资料、有哪些历史处置方案’。"
        "这种从数据到问题的过渡，往往比模型本身更能决定系统是否真的能被使用。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "docker-compose.yml 是另一个“读一遍就知道系统长什么样”的文件。它把 API、Web、Neo4j、Milvus、MinIO 这些组件放在同一张图里，"
        "并且把关键环境变量集中在 .env：例如 Neo4j 密码、Milvus 地址、MinIO 账号、OCR 服务地址、模型目录映射等。"
        "对运维来说，这意味着系统的依赖关系是可视化的；对开发来说，这意味着复现环境成本低，不需要在一台机器上手工装十几个服务。\n\n"
        "你会发现项目对“内网可落地”很在意：比如容器内访问宿主机的 host.docker.internal、"
        "比如把模型目录通过 volume 映射到容器里、比如把镜像导出为 tar 的流程都写成文档与脚本。"
        "这些都不是花活，它们决定了系统能不能真的被部署到业务现场。"
    )
    _p(document, t)
    all_texts.append(t)

    document.add_heading("第五章 部署与运维：能跑起来，也要长期跑得住", level=1)
    t = (
        "项目把依赖尽量装进 docker-compose：API、Web、Neo4j、Milvus、MinIO（以及可选的 MinerU OCR）都在同一套编排里，"
        "健康检查也写进去了。Windows 环境可以用 deploy.ps1 做 start/stop/status/logs/clean；也可以用 docker compose up -d 走标准流程。\n\n"
        "第一次启动通常会卡在两件事：一是 .env 没配好（模型、数据库、管理员密码等），二是镜像拉取或构建时间长。"
        "只要 healthcheck 全部变为 healthy，后续运维会轻松很多。\n\n"
        "日常运维时，我们更建议把“状态”变成一种习惯：先看 deploy.ps1 status 或 docker compose ps，"
        "再看日志（deploy.ps1 logs 或 docker compose logs）。别一上来就调参数或重装环境，很多问题其实只是依赖服务没起来、"
        "端口被占用、磁盘满了。\n\n"
        "内网部署的关键不是“把代码拷过去”，而是“把依赖也带过去”。docs/内网部署指南.md 提供了可复现的流程："
        "在联网机器上构建镜像并导出为 tar，通过 U 盘带到内网机器 docker load，再配好 .env 启动服务。"
        "这套方法看起来土，但非常稳：部署动作可审计、可重复，对内网机器的依赖只剩 Docker 与配置文件。\n\n"
        "长期运维上，建议把三类东西纳入备份范围：docker/volumes（Neo4j/Milvus/MinIO 的数据）、saves（配置、任务状态、SQLite 数据库）以及 .env。"
        "备份更建议配合恢复演练：能恢复，才叫备份；否则只是心理安慰。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "如果你把它作为生产系统来对待，还需要关注三个现实问题：资源、边界、习惯。\n\n"
        "资源方面，向量库和图数据库都吃磁盘；本地大模型和 OCR 会吃 GPU；导入任务会吃 CPU。"
        "所以“能跑起来”只是第一步，“能跑得稳”取决于资源规划：磁盘预留、备份窗口、日志轮转、GPU 共享策略。"
        "尤其是 MinerU 这类 OCR 服务，初次构建和模型下载可能非常大，内网环境更建议把镜像提前在联网机器上完整构建好。\n\n"
        "边界方面，开发环境喜欢 reload，生产环境更需要稳定。生产建议关掉 reload，配置合理的 worker 数量，"
        "并把关键配置（.env 与 saves/config/base.yaml）纳入配置管理。系统里很多“看似神秘”的故障，最后都能归结到一件事：配置漂移。\n\n"
        "习惯方面，建议把“看健康检查”变成上线后的第一动作。API、Neo4j、Milvus、MinIO 都有 healthcheck，"
        "当你看到某个组件不健康时，先把问题收敛到那一层，而不是在前端页面里凭感觉找原因。"
        "这套习惯一旦建立，排障效率会提升一个量级。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "内网搬运时还有一些“很朴素但很关键”的细节。\n\n"
        "镜像导出的 tar 往往很大，U 盘如果是 FAT32 会卡在 4GB 单文件限制，所以更建议用 exFAT 或 NTFS。"
        "内网机器如果要跑 OCR 或本地模型，GPU 驱动与容器运行时的兼容性要提前验证；否则你会在最不该花时间的地方（环境搭建）耗掉项目的大半精力。\n\n"
        "这些看起来不像技术亮点，但它们决定了系统能不能被顺利部署到现场，也决定了后续维护成本。"
    )
    _p(document, t)
    all_texts.append(t)

    document.add_heading("（参考）主要组件与端口", level=2)
    _table(
        document,
        headers=["组件", "端口", "说明"],
        rows=[
            ["Web（开发）", "5173", "前端界面（Vite serve）"],
            ["API", "5050", "后端接口与 /docs"],
            ["Neo4j", "7474 / 7687", "图数据库（浏览器/bolt）"],
            ["Milvus", "19530 / 9091", "向量库（服务/health）"],
            ["MinIO", "9000 / 9001", "对象存储（S3/控制台）"],
            ["MinerU OCR（可选）", "30000", "OCR 服务（GPU，profile=all）"],
        ],
    )

    document.add_heading("第六章 使用与验收：怎么证明它真的有用", level=1)
    t = (
        "建议的验收方式不是只看界面，而是用一条完整链路跑通：建库 → 导入资料 → 提问 → 看引用 → 追问 → 看记录。\n\n"
        "资料最好准备三类：一份结构清晰的 Word、一份文字版 PDF，再加一份扫描 PDF 用来验证 OCR。"
        "提问也别太玄，选一些真实会问的句式，比如：\n"
        "“反滤层常见缺陷有哪些，判定依据和处置措施是什么？”\n"
        "“混凝土结构出现裂缝时，先做哪些排查？哪些情况必须立即处置？”\n"
        "“给我一个雨前巡检清单，按风险优先级排序。”\n"
        "再加一个统计类问题触发图谱，比如“典型原因出现最多的前几项是什么”。\n\n"
        "对系统来说，最好的表现不是回答得花哨，而是能把依据拿出来，并且在不知道时敢说“资料里没找到”。"
        "如果系统总是“很自信但没依据”，那往往说明知识库建设或检索策略还需要调整。\n\n"
        "验收时可以刻意做两件“刁钻”的事：第一，问一个资料里确实有答案的问题，看看它能不能把关键条件说全；"
        "第二，问一个资料里没有答案的问题，看看它会不会胡编。一个成熟系统不怕“不会”，它怕的是“乱说”。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "还有一个很“像人用系统”的验收动作：让两个人用同一份资料提问。\n\n"
        "如果系统的答案随着提问措辞变化太大，那说明检索和提示策略还不够稳；"
        "如果系统能在不同问法下稳定找到同一份关键依据，并把结论写成一致的结构，那基本说明底座和流程是可靠的。"
        "工程场景的可用性，很多时候不是看一次答得多好，而是看十次里有多少次能保持稳定。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "如果你需要一个更“像工程验收”的标准，也可以用下面的方式给系统打分：\n\n"
        "第一，看它能不能稳定跑完导入链路。导入失败不是不可接受，关键是失败能不能被看见、能不能被定位、能不能被重试。\n"
        "第二，看它能不能在回答里体现边界。好的回答通常会出现类似“在资料 A 的第 X 段提到……因此在条件 Y 下建议……”这种表达，"
        "而不是“建议如下（但没有任何依据）”。\n"
        "第三，看它能不能经得起追问。用户真实使用时一定会追问细节，系统要能把追问引回到证据，而不是越聊越飘。\n\n"
        "当这三点过了，系统就具备“上线后持续改进”的资格；否则再多功能也容易沦为演示。"
    )
    _p(document, t)
    all_texts.append(t)

    document.add_heading("第七章 演进路线：下一步怎么做更值", level=1)
    t = (
        "如果你希望系统越用越好，而不是越用越乱，有三条建议很实用。\n\n"
        "第一，优先把资料治理做好。标题规范、版本清晰、扫描件可读、术语统一，比换一个更大的模型更能提升答案质量。\n"
        "第二，把能力做成工具，而不是堆提示词。对接巡检工单、设备台账、监测系统时，把数据接入封装成工具交给 Agent 调用，系统才会稳定可维护。\n"
        "第三，用评测把迭代固定下来。挑一批高频问题做成题库，记录期望答案与引用要求，每次更新模型或参数都回归跑一遍，"
        "否则系统会在“感觉变好/感觉变差”之间来回摇摆。\n\n"
        "最后再说一句：这个系统的价值不在于“像不像专家”，而在于“把专家的知识变成可用的资产”。"
        "当它能持续吸收资料、持续被真实问题驱动，它就会从一个演示项目，慢慢变成一个日常工具。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "从经验上看，下一步最值得投入的不是“换更大的模型”，而是把证据展示和知识治理做深。\n\n"
        "证据展示的意思是：让引用更清晰，让用户一眼就能看到答案依据来自哪个文件、哪个片段、甚至哪个页码或标题。"
        "知识治理的意思是：把资料版本、术语一致性、过期内容清理这类事情做成流程，而不是靠人工记忆。\n\n"
        "当证据和治理到位后，再去做更细的权限隔离、更多业务系统的工具接入、更系统的问答评测，才会变得顺理成章。"
        "否则你会发现系统越用越像一个黑箱：大家觉得“有时候挺准，有时候很玄”，但没人说得清哪里出了问题。"
    )
    _p(document, t)
    all_texts.append(t)

    document.add_heading("第八章 建设过程：从可用闭环到可持续演进", level=1)
    t = (
        "很多人会把这类项目想成“接一个大模型接口，再做个页面”。真正做起来才发现，难点并不在接口，而在工程细节："
        "资料怎么进来、怎么清洗、怎么切块、怎么追溯、怎么在内网部署、怎么让运维愿意接手。\n\n"
        "我们在建设过程中有一个很明确的优先级：先做闭环，再做优化。"
        "闭环指的是：资料能导入、检索能跑通、回答能落库、部署能复现。"
        "只要这个闭环是稳的，后续才谈得上更漂亮的引用展示、更细的权限、更系统的评测。\n\n"
        "另一个很关键的取舍是解耦：把知识和模型解耦，把耗时任务和交互解耦。"
        "模型可以换（云端或本地），知识必须留在自己的存储里；导入/OCR/向量化是后台任务，交互层只负责展示进度与结果。"
        "这套取舍让系统既能快速试起来，也能在条件变化时不至于推倒重来。\n\n"
        "最后是落地心态：不要把“像不像专家”当作唯一目标。"
        "在工程场景里，真正决定可信度的往往是两件事：它能不能把依据拿出来、它能不能承认不知道并引导补证据。"
        "当系统做到这一步，它就已经具备价值；当它能持续被真实问题驱动，它就会越来越像一个可靠的同事。"
    )
    _p(document, t)
    all_texts.append(t)

    document.add_heading("附录：配置与快速排查（尽量少，但够用）", level=1)
    t = (
        "配置文件主要分三处：.env 负责连接与密钥；saves/config/base.yaml 负责运行时选择（默认模型、embedding、是否启用 reranker 等）；"
        "src/config/static/models.yaml 负责“系统里能选哪些模型”。如果你把系统放进内网，通常要改动的是模型 base_url、以及 .env 里各类服务地址。\n\n"
        "排查问题时，建议按一个“从外到内”的顺序来：先看服务是不是都 healthy，再看导入任务是不是完成，再看知识库里是不是有内容，"
        "最后才去怀疑模型。很多时候问题并不复杂，比如 embedding 模型没配 key、Milvus 没起来、Neo4j 密码不一致、MinIO 没权限。\n\n"
        "如果是 OCR 相关问题，再检查 MinerU 服务健康与 GPU 资源；如果图谱为空，确认你用的是 LightRAG 类型知识库且 Neo4j 可连。"
        "还有一个经常被忽略的点：磁盘。向量库、图数据库、对象存储都吃磁盘；当磁盘告急时，系统会以各种“看起来无关”的方式报错。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "如果你希望排查更顺一点，可以把“问题分层”当作基本方法：\n\n"
        "当你看到前端打不开，先问自己：是浏览器访问不到端口，还是页面报错？前者优先看容器与端口占用，后者再看前端日志与接口响应。\n"
        "当你看到登录失败，先问自己：是凭证不对，还是账号被锁定，还是数据库不可写？这三种问题的处理方式完全不同。\n"
        "当你看到导入卡住，先问自己：卡在解析、OCR、向量化还是入库？任务中心的进度和 API 日志通常能把问题定位到一步。\n"
        "当你看到检索无结果，先问自己：知识库是否为空、embedding 是否可用、是否选对了知识库；不要第一反应就怀疑模型。\n\n"
        "系统做得再好，也需要一套“像工程一样”的使用和运维方法。你把问题问对了，系统的答案才更容易变得稳定。"
    )
    _p(document, t)
    all_texts.append(t)

    t = (
        "最后补一句“文档维护”的建议。\n\n"
        "系统会持续迭代，文档也应该跟着迭代，但不需要把每个版本都写成厚厚的技术报告。"
        "最值得同步更新的通常只有三类内容：部署步骤是否有变化、关键配置项是否新增或更名、以及一线使用中暴露的高频问题与最佳实践。\n\n"
        "只要这三类内容能跟上，文档就能一直保持“像人写的、拿来就能用”的状态，而不是放在网盘里积灰。\n\n"
        "另外，本文档是由项目内脚本自动生成的，这样做不是为了省事，而是为了保证“文档与项目一起走”。"
        "当你需要更新内容时，只要按项目实际情况改动生成脚本并重新生成即可；目录也可以在 Word 里右键更新。\n\n"
        "如果后续要做更正式的验收或推广，建议把“真实问题清单”和“标准资料包”一起沉淀下来：每次换模型、换配置、换数据，"
        "都用同一套问题和资料跑一遍，这样系统的改进才会是可证明的，而不是靠印象。"
        "这也是把一个“演示可用”的项目，真正变成“日常必用”的工具的关键。"
        "把这套方法坚持下来，系统会越来越稳定，团队也会越来越敢用。"
        "这份文档也会因此越写越轻、越写越准。"
        "到那时，系统就真正融进工作里了，也更可靠、更省心、更扎实。"
    )
    _p(document, t)
    all_texts.append(t)

    return document, _count_chars(all_texts)


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document, chars = build_document()
    document.save(str(OUT_PATH))
    print(f"Generated: {OUT_PATH}")
    print(f"Approx non-whitespace chars: {chars}")


if __name__ == "__main__":
    main()
