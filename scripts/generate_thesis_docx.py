from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

import docx2txt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]


@dataclass(frozen=True)
class ProjectMeta:
    name: str
    version: str
    generated_at: str


def _read_pyproject_version() -> tuple[str, str]:
    pyproject = ROOT / "pyproject.toml"
    if not pyproject.exists():
        return "smart-water", "unknown"
    try:
        import tomllib  # py>=3.11

        data = tomllib.loads(pyproject.read_text(encoding="utf-8"))
        project = data.get("project") or {}
        return str(project.get("name") or "smart-water"), str(project.get("version") or "unknown")
    except Exception:
        return "smart-water", "unknown"


def _find_existing_spec_doc() -> Path | None:
    docs_dir = ROOT / "docs"
    if not docs_dir.exists():
        return None
    candidates = list(docs_dir.glob("*.docx"))
    if not candidates:
        return None
    # Prefer “项目说明书”
    for p in candidates:
        if "项目说明书" in p.name:
            return p
    return candidates[0]


def _safe_read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8", errors="replace")


def _split_paragraphs(text: str) -> list[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    parts = [p.strip() for p in text.split("\n\n")]
    return [p for p in parts if p]


def _slice_between(text: str, start: str, end: str | None) -> str:
    start_idx = text.find(start)
    if start_idx < 0:
        return ""
    if end is None:
        return text[start_idx:]
    end_idx = text.find(end, start_idx + len(start))
    if end_idx < 0:
        return text[start_idx:]
    return text[start_idx:end_idx]


def _extract_spec_sections(spec_text: str) -> dict[str, list[str]]:
    """
    Extract main narrative parts from the existing “项目说明书”，用于复用为论文正文素材。
    """
    sections: dict[str, str] = {
        "preface": _slice_between(spec_text, "前言", "第一章"),
        "ch1": _slice_between(spec_text, "第一章", "第二章"),
        "ch2": _slice_between(spec_text, "第二章", "第三章"),
        "ch3": _slice_between(spec_text, "第三章", "第四章"),
        "ch4": _slice_between(spec_text, "第四章", "第五章"),
        "ch5": _slice_between(spec_text, "第五章", "第六章"),
        "ch6": _slice_between(spec_text, "第六章", "附录"),
        "appendix": _slice_between(spec_text, "附录", None),
    }
    return {k: _split_paragraphs(v) for k, v in sections.items() if v.strip()}


def _set_style_font(style, *, font_name: str, size_pt: int | None = None, bold: bool | None = None) -> None:
    font = style.font
    font.name = font_name
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if style._element is not None:
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), font_name)


def _setup_doc_styles(doc: Document) -> None:
    _set_style_font(doc.styles["Normal"], font_name="宋体", size_pt=12)
    for level in range(1, 4):
        _set_style_font(doc.styles[f"Heading {level}"], font_name="黑体", bold=True)
    _set_style_font(doc.styles["Title"], font_name="黑体", bold=True)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)


def _add_page_break(doc: Document) -> None:
    doc.add_page_break()


def _add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)
    doc.add_paragraph("提示：打开 Word 后，右键目录区域 →“更新域”→“更新整个目录”。")


def _add_code_block(doc: Document, code: str) -> None:
    for line in code.rstrip("\n").splitlines() or [""]:
        p = doc.add_paragraph(line)
        run = p.runs[0] if p.runs else p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(10)


def _iter_repo_tree_digest() -> Iterable[tuple[str, str]]:
    """
    (path, description) pairs for appendix.
    Keep it short and stable (avoid listing thousands of files like node_modules).
    """
    yield ("server/", "FastAPI 后端入口与 API 路由（鉴权、对话、知识库、图谱、任务等）")
    yield ("src/", "核心能力：知识库、模型接入、Agent/工具、存储与公共工具")
    yield ("web/", "前端界面：Vue3 + Ant Design Vue + 图谱/可视化组件")
    yield ("docker-compose.yml", "一键编排：API/Web/Neo4j/Milvus/MinIO/etcd/（可选）OCR")
    yield ("deploy.ps1 / deploy.sh", "启动、停止、查看状态、日志、清理等运维脚本")
    yield ("saves/", "运行态数据：配置、知识库元数据、任务状态、会话等")
    yield ("models/", "本地模型目录（可选，容器内映射到 /models）")
    yield ("docs/", "部署与配置指南、项目说明书等文档")


def main() -> int:
    project_name, project_version = _read_pyproject_version()
    meta = ProjectMeta(
        name=project_name,
        version=project_version,
        generated_at=datetime.now().strftime("%Y-%m-%d"),
    )

    docs_dir = ROOT / "docs"
    docs_dir.mkdir(parents=True, exist_ok=True)
    output_path = docs_dir / "SmartWater_Thesis.docx"

    spec_doc_path = _find_existing_spec_doc()
    spec_text = docx2txt.process(str(spec_doc_path)) if spec_doc_path else ""
    spec_sections = _extract_spec_sections(spec_text) if spec_text else {}

    internal_deploy_md = ""
    for p in docs_dir.glob("*内网*部署*.md"):
        internal_deploy_md = _safe_read_text(p)
        break

    local_llm_md = ""
    for p in docs_dir.glob("*本地*大模型*配置*指南*.md"):
        local_llm_md = _safe_read_text(p)
        break

    doc = Document()
    _setup_doc_styles(doc)

    # ===== Cover =====
    title = doc.add_paragraph("智能水利知识问答系统（Smart Water / Yuxi‑Know）", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("——基于大语言模型的工程知识管理与可追溯问答系统设计与实现")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"版本：v{meta.version}\n").font.size = Pt(12)
    info.add_run(f"日期：{meta.generated_at}\n").font.size = Pt(12)
    info.add_run("作者：__________\n").font.size = Pt(12)
    info.add_run("单位：__________\n").font.size = Pt(12)
    info.add_run("（提示：上方信息为占位符，可按实际项目填写）").font.size = Pt(10)

    _add_page_break(doc)

    # ===== Abstract =====
    doc.add_heading("摘要", level=1)
    abstract = (
        "水利工程资料具有“体量大、格式杂、分散多源、更新频繁、关键信息隐蔽”等特征。"
        "在日常巡检、运行调度、隐患处置与验收追溯等场景中，资料的可检索性与回答的可追溯性直接影响决策效率与风险控制。"
        "传统关键词检索依赖人为选词，难以覆盖同义表达；而仅依赖大语言模型的自由生成又存在事实幻觉与出处缺失风险。"
        "因此，本项目设计并实现了一套面向工程场景的智能水利知识问答系统（Smart Water / Yuxi‑Know），"
        "以“检索增强生成（RAG）+ 证据驱动回答”为核心思想，通过统一文档解析、结构化切块、向量化检索、重排序与图谱检索等手段，"
        "让模型围绕本地资料组织答案并输出可追溯依据。系统后端采用 FastAPI 提供统一 API，前端采用 Vue3 构建交互界面；"
        "存储层引入 Milvus 作为向量库、Neo4j 作为图数据库、MinIO 作为对象存储，并通过任务中心实现导入/OCR/向量化等耗时流程的异步化与可观测。"
        "在部署方面，系统基于 Docker Compose 进行容器化编排，支持内网离线打包迁移，并兼容 Ollama/vLLM 等本地 OpenAI 兼容推理服务。"
        "本文从需求与约束出发，系统性阐述该系统的总体架构、关键原理、实现细节与部署运维方法，并讨论后续可扩展方向。"
    )
    for p in _split_paragraphs(abstract):
        doc.add_paragraph(p)
    doc.add_paragraph("关键词：大语言模型；检索增强生成；向量数据库；知识图谱；OCR；内网部署；水利工程资料治理")

    _add_page_break(doc)

    # ===== TOC =====
    doc.add_heading("目录", level=1)
    _add_toc_field(doc)
    _add_page_break(doc)

    # ===== Body =====
    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph(
        "本章从工程资料管理的现实困境出发，说明研究背景与问题定义，明确系统目标与边界，"
        "并给出本文结构安排。与“纯生成式对话”不同，本项目强调“证据优先、可追溯优先、可治理优先”，"
        "将大模型定位为“表达与组织”的助手，而不是替代专家判断的最终裁决者。"
    )
    if preface := spec_sections.get("preface"):
        doc.add_heading("1.1 项目缘起与场景叙述", level=2)
        for p in preface:
            doc.add_paragraph(p)
    if ch1 := spec_sections.get("ch1"):
        doc.add_heading("1.2 痛点分析与系统边界", level=2)
        for p in ch1:
            doc.add_paragraph(p)
    doc.add_heading("1.3 本文贡献与结构", level=2)
    doc.add_paragraph(
        "本文的主要工作包括："
        "（1）提出一套面向水利工程资料的“RAG + 图谱检索 + 异步任务”的系统化方案；"
        "（2）实现可落地的多源文档导入与解析流水线，覆盖常见办公文档与扫描件；"
        "（3）实现可配置的模型接入与本地/云端兼容部署；"
        "（4）围绕内网场景提供镜像打包与一键部署脚本，降低交付成本。"
    )

    doc.add_heading("第二章 理论基础与关键原理", level=1)
    doc.add_paragraph(
        "本章介绍系统背后的核心原理：向量表示与相似度检索、RAG 的证据驱动回答、重排序的精排思想、"
        "以及知识图谱在结构化关系查询中的价值。对于工程场景而言，原理本身并不新，关键在于如何将其“工程化”："
        "把数据治理、任务编排、可追溯性与可运维性一起纳入设计。"
    )
    doc.add_heading("2.1 向量表示与相似度检索", level=2)
    doc.add_paragraph(
        "Embedding 模型将文本片段映射为高维向量，使得语义相近的内容在向量空间中距离更近。"
        "检索时常用余弦相似度：cos(q, d) = (q·d) / (|q||d|)。"
        "工程资料的挑战在于：长文档结构复杂、同义表达丰富、关键条件分散，因此必须在“切块策略、召回策略、元数据治理”上做工程化取舍。"
    )
    doc.add_heading("2.2 检索增强生成（RAG）", level=2)
    doc.add_paragraph(
        "RAG 的核心思想是：先检索，再生成。系统将检索到的片段作为证据输入模型，要求模型在证据范围内组织答案，"
        "并尽可能输出引用来源，从而降低幻觉风险并提升可审计性。"
        "与传统搜索不同，RAG 既需要“召回”，也需要“可读性”：返回的片段要能被人理解与核对。"
    )
    doc.add_heading("2.3 重排序（Re-rank）", level=2)
    doc.add_paragraph(
        "粗召回阶段通常追求速度与覆盖，难免带入噪声。重排序模型（如 bge-reranker 系列）"
        "可对候选片段进行精排，提升“最相关片段”出现在前几位的概率。"
        "在实现上，本项目将重排序作为可配置能力（见 `src/config/app.py`），以兼容不同算力与成本约束。"
    )
    doc.add_heading("2.4 知识图谱与图检索（LightRAG 思路）", level=2)
    doc.add_paragraph(
        "当问题涉及实体关系、因果链条、治理措施关联等结构化信息时，单纯的向量检索可能缺少“关系路径”。"
        "本项目提供 LightRAG 类型知识库：在向量库之外，引入 Neo4j 存储实体与关系，支持向量相似度召回实体、"
        "再以多跳扩展形成可解释的关系子图，从而在“解释性、结构化统计、关联推理”上更贴近工程语境。"
    )
    if ch2 := spec_sections.get("ch2"):
        doc.add_heading("2.5 系统工作流程叙述（结合本项目）", level=2)
        for p in ch2:
            doc.add_paragraph(p)

    doc.add_heading("第三章 需求分析与总体约束", level=1)
    doc.add_heading("3.1 角色与典型使用场景", level=2)
    doc.add_paragraph(
        "系统面向的主要角色包括：普通使用者（提问与查看依据）、资料管理员（导入与维护知识库）、"
        "系统管理员（账号与部署运维）。典型场景包括：巡检隐患查询、缺陷与处置措施检索、规程条款定位、"
        "历史案例复盘、跨文档对比与解释。"
    )
    doc.add_heading("3.2 功能性需求", level=2)
    doc.add_paragraph(
        "围绕上述场景，系统需提供：知识库创建与管理、文档/URL 导入、异步任务与进度展示、对话与多智能体调用、"
        "图谱查询与可视化、权限控制与审计、数据导出与迁移等能力。相关 API 路由集中在 `server/routers/` 目录。"
    )
    doc.add_heading("3.3 非功能性需求", level=2)
    doc.add_paragraph(
        "工程化约束主要体现在：内网可部署与可离线交付；可观测（日志、健康检查、任务进度）；"
        "可靠性（服务恢复、数据持久化）；安全性（登录与权限、必要的限流与锁定策略）；"
        "以及可扩展（模型可替换、存储可替换、知识库类型可扩展）。"
    )

    doc.add_heading("第四章 系统总体设计", level=1)
    doc.add_heading("4.1 总体架构", level=2)
    doc.add_paragraph(
        "系统采用典型的前后端分离与多存储组合架构：前端负责交互与可视化，后端提供 API 与业务编排；"
        "向量库与图数据库分别承担“语义召回”与“关系表达”，对象存储用于原文件与图片等二进制资产。"
    )
    doc.add_paragraph("表 4‑1 给出各组件的职责划分。")
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "组件"
    hdr_cells[1].text = "职责与价值"
    rows = [
        ("Web（Vue3）", "对话、知识库管理、任务中心、图谱与统计可视化"),
        ("API（FastAPI）", "统一业务入口：鉴权、知识库/图谱编排、模型调用、流式输出"),
        ("Milvus", "向量检索：文档片段与候选证据召回"),
        ("Neo4j", "图谱存储与关系查询：实体‑关系‑多跳扩展、结构化统计"),
        ("MinIO", "对象存储：原文档/图片等资源统一保存与访问"),
        ("MinerU（可选）", "OCR 与版面解析：提升扫描件/图片的可用性"),
    ]
    for comp, desc in rows:
        r = table.add_row().cells
        r[0].text = comp
        r[1].text = desc

    doc.add_heading("4.2 数据流与任务流", level=2)
    doc.add_paragraph(
        "系统将“交互”与“耗时处理”分离：用户导入资料后，后台任务负责解析、OCR、切块、向量化、入库；"
        "对话时则以“检索 → 组装证据 → 生成回答”的在线流程为主。"
        "任务系统实现见 `server/services/tasker.py`，可提供进度、结果与取消能力。"
    )
    doc.add_paragraph("图 4‑1 以文本形式概括关键流程：")
    _add_code_block(
        doc,
        """
用户上传/选择资料
        ↓
文档解析（Word/PDF/Excel/图片… → Markdown/文本）
        ↓
切块（chunk_size/overlap 或 QA 切分）
        ↓
Embedding 向量化
        ↓
写入知识库（Milvus / LightRAG：Milvus + Neo4j）
        ↓
用户提问 → 检索召回 →（可选）重排序 → 证据组装 → LLM 生成（附出处）
""".strip("\n"),
    )

    doc.add_heading("第五章 系统实现", level=1)
    doc.add_heading("5.1 后端实现：路由、鉴权与流式对话", level=2)
    doc.add_paragraph(
        "后端入口为 `server/main.py`，采用 FastAPI 构建，并注册鉴权、对话、知识库、图谱、任务等路由（见 `server/routers/__init__.py`）。"
        "系统支持登录失败限流与锁定、JWT 鉴权，以及按路径区分公开/受保护 API。"
        "对话接口支持流式输出，以便在长回答与工具调用场景下提升交互体验。"
    )
    doc.add_heading("5.2 知识库实现：Milvus 与 LightRAG 两条路径", level=2)
    doc.add_paragraph(
        "知识库模块位于 `src/knowledge/`。其中 MilvusKB 侧重“生产级向量检索”，"
        "将文档片段与元数据写入 Milvus 集合；LightRagKB 则在此基础上引入实体关系抽取与 Neo4j 存储，"
        "支持更结构化的关联查询。两类知识库通过工厂注册与管理器统一调度（见 `src/knowledge/__init__.py`）。"
    )
    doc.add_heading("5.3 文档解析与 OCR", level=2)
    doc.add_paragraph(
        "文档导入统一走“转 Markdown/文本 → 切块”链路。对于扫描 PDF、图片等非结构化输入，"
        "系统可接入 OCR 服务（MinerU / PaddleX），在尽量保留结构的前提下提取可检索文本。"
        "解析与 OCR 的工程价值不在于“把字识出来”这么简单，而在于让后续的切块更接近文档语义边界。"
    )
    doc.add_heading("5.4 对话 Agent 与工具调用", level=2)
    doc.add_paragraph(
        "系统采用 LangGraph 组织 Agent 流程（见 `src/agents/chatbot/graph.py`），"
        "将检索、图谱查询、统计类问题处理等能力以工具形式接入，并把工具调用参数与结果持久化到会话存储中。"
        "这种设计的意义在于：回答不仅是文本，还能追溯“用了哪些工具、查了哪些证据”。"
    )
    if ch4 := spec_sections.get("ch4"):
        doc.add_heading("5.5 实现细节补充（项目说明书摘录）", level=2)
        for p in ch4:
            doc.add_paragraph(p)
    if ch5 := spec_sections.get("ch5"):
        for p in ch5:
            doc.add_paragraph(p)

    doc.add_heading("第六章 部署与运维", level=1)
    doc.add_heading("6.1 容器化编排与一键启动", level=2)
    doc.add_paragraph(
        "系统通过 `docker-compose.yml` 编排 API/Web/Neo4j/Milvus/MinIO 等服务，"
        "并提供 `deploy.ps1`/`deploy.sh` 脚本封装常用运维动作（start/stop/status/logs/clean）。"
        "在内网交付时，可先在联网环境构建镜像并导出为 tar，再在目标机加载并启动。"
    )
    doc.add_heading("6.2 内网离线部署要点", level=2)
    if internal_deploy_md:
        doc.add_paragraph(
            "项目提供了内网部署指南（`docs/内网部署指南.md`）。本文仅总结关键步骤："
            "（1）在联网机拉取/构建镜像；（2）导出镜像 tar 并拷贝项目代码；（3）内网机 load 镜像、配置 `.env`、启动服务；"
            "（4）通过健康检查确认各容器正常。"
        )
        doc.add_paragraph("关键命令（节选）：")
        _add_code_block(
            doc,
            r"""
# 联网机：构建并导出镜像（PowerShell）
.\docker\save_docker_images.ps1

# 内网机：加载镜像
docker load -i docker_images_backup\\smart_water_images_*.tar

# 内网机：启动
.\deploy.ps1 start
""".strip("\n"),
        )
    else:
        doc.add_paragraph("未在 `docs/` 中发现内网部署指南文件，但项目已提供镜像打包与一键部署脚本。")

    doc.add_heading("6.3 本地大模型接入与配置", level=2)
    if local_llm_md:
        doc.add_paragraph(
            "系统支持对接本地 OpenAI 兼容接口（如 Ollama、vLLM），也支持云端模型供应商。"
            "配置入口包括：`.env`（连接与密钥）、`saves/config/base.yaml`（默认模型选择）、"
            "`src/config/static/models.yaml`（可选模型列表）。"
        )
        doc.add_paragraph("本地模型配置要点（节选）：")
        _add_code_block(
            doc,
            """
# .env 示例：指向内网 vLLM（OpenAI 兼容）
OPENAI_API_BASE=http://<LLM_HOST>:8000/v1
OPENAI_API_KEY=no-key-needed
""".strip("\n"),
        )
    else:
        doc.add_paragraph("未在 `docs/` 中发现本地大模型配置指南文件，但系统具备 OpenAI 兼容接口接入能力。")

    if appendix := spec_sections.get("appendix"):
        doc.add_heading("6.4 运维经验与排查方法（项目说明书摘录）", level=2)
        for p in appendix:
            doc.add_paragraph(p)

    doc.add_heading("第七章 测试、评估与治理方法", level=1)
    doc.add_heading("7.1 评估维度", level=2)
    doc.add_paragraph(
        "在工程场景中，系统效果不应只用“像不像专家”衡量，更重要的是："
        "（1）可追溯：是否能给出出处与证据；（2）可复核：证据是否足以让人核对结论；"
        "（3）可用性：响应时间与交互是否可接受；（4）可治理：资料新增/变更后是否可增量更新与回滚。"
    )
    doc.add_heading("7.2 基准问题集与回归方法", level=2)
    doc.add_paragraph(
        "建议为系统建立“标准资料包 + 基准问题集”的回归方法：每次更换模型、切块参数或知识库类型，"
        "都用同一套问题验证回答与出处是否稳定。该方法能够将迭代从“凭印象”变为“可证明”。"
    )
    doc.add_heading("7.3 数据治理建议", level=2)
    doc.add_paragraph(
        "系统落地的关键往往不是模型，而是资料治理：统一命名与版本、完善元数据（来源、时间、适用范围）、"
        "对敏感信息做分级与脱敏、建立“哪些资料必须入库”的清单。RAG 系统会把资料不一致暴露出来，"
        "这既是挑战，也是推进治理的抓手。"
    )

    doc.add_heading("第八章 总结与展望", level=1)
    doc.add_paragraph(
        "本文围绕 Smart Water / Yuxi‑Know 项目，给出一套面向工程资料的可落地问答系统方案。"
        "系统通过向量检索与图谱检索结合，强化回答的证据基础；通过异步任务与容器化部署，降低导入与交付门槛；"
        "通过可配置模型接入，适应内网与多供应商环境。"
    )
    doc.add_paragraph(
        "后续工作可从以下方向扩展："
        "（1）更精细的权限与审计（按知识库/文档/字段级别）；"
        "（2）增量索引与版本化知识库；"
        "（3）更强的多模态解析（图表、公式、图纸）；"
        "（4）面向业务系统的接口集成与流程自动化；"
        "（5）基于真实问题集的系统化评测与持续回归。"
    )

    # ===== References & Appendix =====
    doc.add_heading("参考文献", level=1)
    refs = [
        "[1] Lewis P, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS, 2020.",
        "[2] Vaswani A, et al. Attention Is All You Need. NeurIPS, 2017.",
        "[3] Milvus Documentation. Vector Database for Scalable Similarity Search.",
        "[4] Neo4j Documentation. Graph Database & Cypher Query Language.",
        "[5] LangChain / LangGraph Documentation. Agentic Workflow Orchestration.",
    ]
    for r in refs:
        doc.add_paragraph(r)

    doc.add_heading("附录A 项目目录结构与职责划分", level=1)
    for path, desc in _iter_repo_tree_digest():
        doc.add_paragraph(f"{path}：{desc}")

    doc.add_heading("附录B 关键配置入口清单", level=1)
    doc.add_paragraph("`.env`：服务连接地址与密钥（模型、Neo4j、Milvus、MinIO 等）。")
    doc.add_paragraph("`saves/config/base.yaml`：运行时默认选择（默认模型、Embedding、是否启用 Re-ranker 等）。")
    doc.add_paragraph("`src/config/static/models.yaml`：系统中可选模型列表与供应商配置。")
    doc.add_paragraph("`docker-compose.yml`：容器编排与端口映射。")

    if internal_deploy_md:
        doc.add_heading("附录C 内网部署指南（原文节选）", level=1)
        snippet = "\n".join(internal_deploy_md.splitlines()[:120])
        _add_code_block(doc, snippet)

    # Persist
    doc.save(str(output_path))

    # Lightweight sanity report (stdout)
    try:
        extracted = docx2txt.process(str(output_path))
        char_count = len(extracted.replace("\n", "").strip())
    except Exception:
        char_count = -1

    print(f"Generated: {output_path}")
    print(f"CharacterCount(rough): {char_count}")
    if spec_doc_path:
        print(f"BaseMaterial: {spec_doc_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
